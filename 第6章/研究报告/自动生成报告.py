# -*- coding: utf-8 -*-
"""
Created on Tue Aug  4 20:55:52 2020

@author: Administrator
"""
name="四川公司"
from docx import Document
myDocx=Document(r'H:\示例\第6章\研究报告\模板.docx')
tb=myDocx.tables[0].cell(0,1).tables[0]
xlFile='H:\\示例\\第6章\\研究报告\\'+name+'数据库.xls'
import pandas as pd
df=pd.read_excel(xlFile,['基本资料','财务指标'],header=None)
sheet0=df['基本资料']
sheet1=df['财务指标']
tb.cell(1,0).paragraphs[0].text=''
run=tb.cell(1,0).paragraphs[0].add_run(sheet0.iat[0,1])
run.font.name = u'微软雅黑'
from docx.oxml.ns import qn  
run._element.rPr.rFonts.set(qn('w:eastAsia'),u'微软雅黑')
from docx.shared import Pt
run.font.size=Pt(25)
tb.cell(4,0).paragraphs[1].text=sheet0.iat[17,1]
tb.cell(4,0).paragraphs[3].text=sheet0.iat[18,1]
myDocx.styles['Normal'].font.name = u'宋体'
myDocx.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
myDocx.styles['Normal'].font.size=Pt(11)
from docx.enum.text import WD_ALIGN_PARAGRAPH
for i in range(0,17):
    tb.cell(4,1).tables[0].cell(i+1,1).text=sheet0.iat[i,1]
    prg=tb.cell(4,1).tables[0].cell(i+1,1).paragraphs[0]
    prg.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.RIGHT
myDocx.styles['Normal'].font.size=Pt(9)
for r in range(2,10):
    for c in range(1,6):
        value=sheet1.iat[r-1,c]
        try:
            value=str(round(float(value),2))
        except:
            pass
        tb.cell(4,0).tables[0].cell(r,c).text=value
import numpy as np
data=sheet1.loc[9][1:6].astype(np.float)
import matplotlib.pyplot as plt
plt.figure(figsize=(7, 4))
title=sheet1.iloc[0].tolist()
title.pop(0)
fig = plt.bar([1,2,3,4,5],data,tick_label=title,color='black')
plt.tick_params(labelsize=20)
def add_datalabels(bars):
    for bar in bars:
        height = bar.get_height()
        x=bar.get_x() + bar.get_width()/2
        y=height+0.01*height
        plt.text(x,y,'%.0f'%height,ha='center',va='bottom',fontsize=30)
        bar.set_edgecolor('white')
add_datalabels(fig)
filePic=r'H:\示例\第6章\研究报告\pic_2.png'
plt.savefig(filePic)
paragraph=tb.cell(5,1).tables[0].cell(1,0).paragraphs[0]
paragraph.text=''
from docx.shared import Inches
paragraph.add_run().add_picture(filePic,width=Inches(2.8))
myDocx.save('H:\\示例\\第6章\\研究报告\\'+name+'分析报告.docx')


